"""Gate for the draft -> Word converter (local/draft2docx.mjs).

Converts a representative TestPlanGen draft (banner comment, WARNING
alert, verifier IMPORTANT block comment, task lists, GFM tables,
bold labels, the Issue Trace addendum) and reads the .docx back with
python-docx, asserting the conversion contract:

  1. the package opens (python-docx = the same OOXML stack Word uses
     for validation-grade parsing)
  2. headings land as real Heading 1/2/3 styles, in draft order —
     the Word navigation pane shows the plan structure
  3. GFM tables become Word tables with the header row bold and every
     cell's text intact (Overview, Source Case Sweep, Coverage Map,
     Issue Trace)
  4. task-list items render as checkbox glyphs in List Paragraph
     style, numbering preserved verbatim
  5. alert blocks render as a bold label + Quote-styled body;
     **bold** spans become bold runs (Expected Result / Trace labels)
  6. HTML comments (the machine banner, the verify stamp) are DROPPED
     — no machine plumbing reaches the document of record
  7. wrapped prose lines join into single paragraphs
  8. CLI contract: sibling-name default, -o for a single input,
     multi-input conversion, usage errors exit nonzero

Needs python-docx (review/harness/requirements.txt — the CI
full-format job installs it). Usage: python3 check_draft2docx.py
"""
import os
import subprocess
import sys
import tempfile

HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.dirname(os.path.dirname(HERE))
JOB = os.path.join(REPO, "local", "draft2docx.mjs")

PASS = []
FAIL = []


def check(name, cond, detail=""):
    (PASS if cond else FAIL).append(name)
    mark = "ok  " if cond else "FAIL"
    print(f"  {mark} {name}" + ("" if cond else f"  <- {detail}"))


DRAFT = """<!-- machine-generated test-plan draft — TestPlanGen prompt v1.7 · local/testplangen.mjs v1.3 · provider anthropic -->
> [!WARNING]
> **DRAFT — machine-generated, unreviewed.** Generated 2026-09-04T00:00:00Z from user story doc 12 — "Route Merge". Source sidecar: <https://example/sc.md>
> Review every case and resolve all [VERIFY] items before use.

<!-- verify: 1 finding(s) — lib/draftlint.mjs, prompt v1.7 contract + grounding -->
> [!IMPORTANT]
> Draft verifier: 1 finding(s) — review these first:
> - TC-N1 carries a **Trace:** line

# Test Plan — Route Merge

## Overview

| Surface | Target release | PE |
| --- | --- | --- |
| Pro | 3.8 | Claire Wang |

Verifies measure-preserving merge of two routes in ArcGIS Pro,
covering the workflow the story enumerates.

## Setup / Prerequisites
- [ ] 1. LRS network with two mergeable routes. [VERIFY: minimum network configuration]
- [x] 2. Two Pro sessions signed in.

## Positive Tests

### TC-P1 — Merge preserves measures
**Steps:**
- [ ] 1. Run Merge Routes on route A and route B.
- [ ] 2. Inspect the measures on the merged route.

**Expected Result:** The merged route keeps the source measures unchanged.

**Trace:** "the merge must preserve measures" — story requirement.

## Negative Tests

> [!CAUTION]
> A pass below is the described denial or error — never the edit
> succeeding.

### TC-N1 — Merge denied on locked route
**Steps:**
- [ ] 1. As user B, attempt Merge Routes on a locked route.

**Expected Result:** The merge is denied with a lock conflict.

**Trace:** exemplar pattern — multi-user denial case.

## Open Questions
- [ ] [VERIFY: minimum network configuration for setup]

## Coverage Map

| # | Requirement (source) | Covered by |
| --- | --- | --- |
| 1 | "the merge must preserve measures" (requirement) | TC-P1 |
| 2 | denial on locked routes (conflict statement) | TC-N1 |

## Issue Trace

_Deterministic addendum — minted by local/testplangen.mjs from the Doc IDs and Issue Refs lists, not by the model. Cross-check against devtopia during the review pass._

| Issue | Title (Issue Refs) | Schedule status | Found via |
| --- | --- | --- | --- |
| ArcGISPro/ps-location-referencing#4855 | Route merge epic | Iteration 2 · Dev=Completed | sidecar |
"""


def main():
    try:
        import docx
    except ImportError:
        print("SKIP: python-docx not installed (pip install -r requirements.txt)")
        sys.exit(0)

    tmp = tempfile.mkdtemp(prefix="draft2docx-gate-")
    md1 = os.path.join(tmp, "TestPlanDraft__doc12__20260904-000000.md")
    with open(md1, "w", encoding="utf-8") as f:
        f.write(DRAFT)

    # ---- CLI: sibling-name default -----------------------------------
    r = subprocess.run(["node", JOB, md1], capture_output=True, text=True, cwd=REPO)
    out1 = md1[:-3] + ".docx"
    check("converts with the sibling name", r.returncode == 0 and os.path.exists(out1),
          r.stdout + r.stderr)
    check("reports paragraph/table counts", "paragraphs" in r.stdout and "tables" in r.stdout,
          r.stdout)

    d = docx.Document(out1)
    paras = d.paragraphs
    texts = [p.text for p in paras]
    styles = [p.style.name for p in paras]

    # 2 — headings, in order
    heads = [(p.style.name, p.text) for p in paras if p.style.name.startswith("Heading")]
    check("H1 title as Heading 1",
          ("Heading 1", "Test Plan — Route Merge") in heads, str(heads[:3]))
    h2 = [t for s, t in heads if s == "Heading 2"]
    check("H2 sections in draft order",
          h2 == ["Overview", "Setup / Prerequisites", "Positive Tests",
                 "Negative Tests", "Open Questions", "Coverage Map", "Issue Trace"],
          str(h2))
    check("TC cases as Heading 3",
          [t for s, t in heads if s == "Heading 3"] ==
          ["TC-P1 — Merge preserves measures", "TC-N1 — Merge denied on locked route"],
          str(heads))

    # 3 — tables (Overview, Coverage Map, Issue Trace)
    check("all three tables arrive", len(d.tables) == 3, str(len(d.tables)))
    check("Overview cells intact",
          [c.text for c in d.tables[0].rows[1].cells] == ["Pro", "3.8", "Claire Wang"],
          str([c.text for c in d.tables[0].rows[0].cells]))
    check("table header row bold",
          all(r.bold for c in d.tables[0].rows[0].cells
              for p in c.paragraphs for r in p.runs), "")
    cm = d.tables[1]
    check("Coverage Map rows survive",
          "the merge must preserve measures" in cm.rows[1].cells[1].text
          and cm.rows[1].cells[2].text == "TC-P1", str([c.text for c in cm.rows[1].cells]))
    check("Issue Trace row survives",
          "ArcGISPro/ps-location-referencing#4855" in d.tables[2].rows[1].cells[0].text
          and "Iteration 2" in d.tables[2].rows[1].cells[2].text,
          str([c.text for c in d.tables[2].rows[1].cells]))

    # 4 — task lists
    boxes = [t for t, s in zip(texts, styles) if s == "List Paragraph" and t[:1] in "☐☑"]
    check("unchecked items as U+2610 with numbering verbatim",
          any(t.startswith("☐ 1. LRS network") for t in boxes), str(boxes[:3]))
    check("checked items as U+2611",
          any(t.startswith("☑ 2. Two Pro sessions") for t in boxes), str(boxes))

    # 5 — alerts + bold
    labels = [t for t, s in zip(texts, styles) if s == "Alert Label"]
    check("alert labels rendered bold-styled",
          labels == ["WARNING", "IMPORTANT", "CAUTION"], str(labels))
    check("alert body in Quote style",
          any(s == "Quote" and "never the edit succeeding" in t
              for t, s in zip(texts, styles)), "")
    check("bold label runs",
          any(r.bold and r.text == "Expected Result:" for p in paras for r in p.runs), "")
    check("unbalanced ** stays literal (none present here)",
          not any("**" in t for t in texts), str([t for t in texts if "**" in t][:2]))

    # 6 — comments dropped
    check("HTML comments dropped",
          not any("machine-generated test-plan draft" in t or "verify: 1 finding" in t
                  for t in texts), "")

    # 7 — wrapped prose joins
    check("wrapped prose joins into one paragraph",
          any(t == "Verifies measure-preserving merge of two routes in ArcGIS "
                   "Pro, covering the workflow the story enumerates." for t in texts),
          str([t for t in texts if t.startswith("Verifies")]))

    # metadata
    check("core title from the H1", d.core_properties.title == "Test Plan — Route Merge", "")

    # ---- CLI: -o, multi-input, usage ---------------------------------
    out2 = os.path.join(tmp, "renamed.docx")
    r = subprocess.run(["node", JOB, md1, "-o", out2], capture_output=True, text=True, cwd=REPO)
    check("-o names a single output", r.returncode == 0 and os.path.exists(out2), r.stderr)
    md2 = os.path.join(tmp, "second.md")
    with open(md2, "w", encoding="utf-8") as f:
        f.write("# Second\n\n## Overview\n\nBody.\n")
    r = subprocess.run(["node", JOB, md1, md2], capture_output=True, text=True, cwd=REPO)
    check("multi-input converts each beside its source",
          r.returncode == 0 and os.path.exists(os.path.join(tmp, "second.docx")), r.stderr)
    check("usage: no inputs exits nonzero",
          subprocess.run(["node", JOB], capture_output=True).returncode != 0, "")
    check("usage: -o with several inputs exits nonzero",
          subprocess.run(["node", JOB, md1, md2, "-o", out2],
                         capture_output=True).returncode != 0, "")

    print(f"\n{len(PASS)} passed, {len(FAIL)} failed")
    if FAIL:
        print("FAILED: " + ", ".join(FAIL))
        sys.exit(1)
    print("RESULT: PASS")


if __name__ == "__main__":
    main()
